"""
Сравнение эффективности поиска: PostgreSQL+pgvector vs Qdrant vs Qdrant+квантование
vs полная production-инфраструктура проекта (Qdrant + квантование + тюнинг HNSW + кэш).

Это скрипт для РЕАЛЬНОГО прогона на вашей инфраструктуре. Он не подделывает цифры —
если Postgres или Qdrant недоступны, соответствующий вариант помечается как
"unavailable" в отчёте, а не подставляется вручную.

Запуск:
    python -m app.scripts.compare_search_methods \
        --database-url postgresql://postgres:postgres@localhost:5432/products \
        --qdrant-host localhost \
        --vectors 50000 --queries 300

Результат:
    search_comparison_report.json  — сырые метрики (для CI / трекинга регрессий)
    search_comparison_report.md    — читаемый отчёт
    search_comparison_report.docx  — оформленный документ для демонстрации заказчику

Методология:
1. Один и тот же набор случайных нормализованных векторов (косинусная метрика)
   и один и тот же набор запросов используется для ВСЕХ 4 вариантов — это
   единственный способ сравнивать честно.
2. Ground truth (для recall@10) считается ОДИН раз через brute-force в numpy,
   независимо от какой-либо БД — так что recall не зависит от того, что каждый
   движок сам считает "точным" поиском.
3. Каждый вариант запускается на одинаковом K=10, с прогревом (первые 10% запросов
   отбрасываются из статистики латентности, чтобы не мерить холодный старт JIT/кэшей
   на уровне ОС).
"""

from __future__ import annotations

import argparse
import asyncio
import json
import logging
import statistics
import time
import uuid
from dataclasses import dataclass, asdict
from pathlib import Path

import numpy as np

logging.basicConfig(level=logging.INFO, format="%(asctime)s | %(levelname)-8s | %(message)s")
logger = logging.getLogger("compare_search")

DIM = 384
K = 10
WARMUP_FRACTION = 0.1

# Параметры HNSW — одинаковые везде, где применимо, чтобы сравнение было честным
# (варианты отличаются наличием квантования и кэша, а не разной настройкой HNSW).
HNSW_M = 32
HNSW_EF_CONSTRUCT = 200
HNSW_EF_SEARCH = 128

# Предполагаемая доля повторных запросов в реальном трафике (для варианта 4,
# "полная инфраструктура" — там есть кэш результатов). Это ЯВНОЕ предположение,
# а не измеренная величина — в отчёте помечается как assumption, и в проде
# нужно заменить на реальный cache_hit_rate из /health после недели трафика.
ASSUMED_CACHE_HIT_RATE = 0.30


@dataclass
class VariantResult:
    variant: str
    description: str
    available: bool
    recall_at_10: float | None = None
    p50_ms: float | None = None
    p95_ms: float | None = None
    p99_ms: float | None = None
    qps: float | None = None
    notes: str = ""


def _gen_vectors(n: int, dim: int, seed: int) -> np.ndarray:
    rng = np.random.default_rng(seed)
    v = rng.normal(size=(n, dim)).astype(np.float32)
    v /= np.linalg.norm(v, axis=1, keepdims=True)
    return v


def _brute_force_ground_truth(vectors: np.ndarray, queries: np.ndarray, k: int) -> list[list[int]]:
    """Единый, независимый от БД, ground truth для recall — косинусное сходство в numpy."""
    logger.info("Computing brute-force ground truth for %d queries over %d vectors", len(queries), len(vectors))
    sims = queries @ vectors.T  # векторы уже нормализованы -> dot product == cosine sim
    top_k = np.argsort(-sims, axis=1)[:, :k]
    return top_k.tolist()


def _recall(ground_truth: list[list[int]], predicted_indices: list[list[int]]) -> float:
    scores = []
    for gt, pred in zip(ground_truth, predicted_indices):
        if not gt:
            continue
        scores.append(len(set(gt) & set(pred)) / len(gt))
    return round(statistics.mean(scores), 4) if scores else 0.0


def _trimmed_percentiles(latencies: list[float], warmup_fraction: float) -> tuple[float, float, float, float]:
    """Отбрасывает warmup-долю запросов, возвращает p50/p95/p99 + qps по оставшимся."""
    n_warmup = int(len(latencies) * warmup_fraction)
    trimmed = latencies[n_warmup:] or latencies
    s = sorted(trimmed)
    n = len(s)

    def p(pct: float) -> float:
        return round(s[min(int(n * pct / 100), n - 1)], 3)

    qps = round(1000 / statistics.mean(trimmed), 2) if trimmed else 0.0
    return p(50), p(95), p(99), qps


# ----------------------------------------------------------------------------
# Вариант 1: чистый PostgreSQL + pgvector (HNSW-индекс, те же параметры, что у Qdrant)
# ----------------------------------------------------------------------------

async def bench_postgres_pgvector(
    database_url: str,
    vectors: np.ndarray,
    queries: np.ndarray,
    ground_truth: list[list[int]],
) -> VariantResult:
    variant = "1. PostgreSQL + pgvector"
    description = (
        f"HNSW-индекс pgvector (m={HNSW_M}, ef_construction={HNSW_EF_CONSTRUCT}), "
        f"ef_search={HNSW_EF_SEARCH}, cosine distance. Без Qdrant, без квантования."
    )
    try:
        import asyncpg
    except ImportError:
        return VariantResult(variant, description, available=False, notes="asyncpg не установлен")

    table = f"bench_pgvector_{uuid.uuid4().hex[:8]}"
    try:
        conn = await asyncpg.connect(dsn=database_url, timeout=10)
    except Exception as exc:
        return VariantResult(variant, description, available=False, notes=f"Postgres недоступен: {exc}")

    try:
        await conn.execute("CREATE EXTENSION IF NOT EXISTS vector")
        await conn.execute(f"CREATE TABLE {table} (id INT PRIMARY KEY, embedding vector({vectors.shape[1]}))")

        logger.info("[pgvector] Inserting %d vectors", len(vectors))
        records = [(i, "[" + ",".join(f"{x:.6f}" for x in vec) + "]") for i, vec in enumerate(vectors)]
        await conn.executemany(f"INSERT INTO {table} (id, embedding) VALUES ($1, $2::vector)", records)

        try:
            await conn.execute(
                f"CREATE INDEX ON {table} USING hnsw (embedding vector_cosine_ops) "
                f"WITH (m = {HNSW_M}, ef_construction = {HNSW_EF_CONSTRUCT})"
            )
            index_note = "HNSW"
        except Exception as exc:
            # Старые версии pgvector (<0.5) не умеют HNSW — fallback на IVFFlat,
            # честно отметим это в отчёте, т.к. это другой алгоритм с другими характеристиками.
            logger.warning("[pgvector] HNSW index unavailable (%s), falling back to IVFFlat", exc)
            lists = max(int(len(vectors) ** 0.5), 10)
            await conn.execute(
                f"CREATE INDEX ON {table} USING ivfflat (embedding vector_cosine_ops) WITH (lists = {lists})"
            )
            index_note = f"IVFFlat (lists={lists}) — pgvector <0.5, HNSW недоступен"

        await conn.execute(f"ANALYZE {table}")
        try:
            await conn.execute(f"SET hnsw.ef_search = {HNSW_EF_SEARCH}")
        except Exception:
            pass  # ivfflat fallback path — параметр не применим

        latencies: list[float] = []
        predicted: list[list[int]] = []
        for q in queries:
            qvec = "[" + ",".join(f"{x:.6f}" for x in q) + "]"
            t0 = time.perf_counter()
            rows = await conn.fetch(
                f"SELECT id FROM {table} ORDER BY embedding <=> $1::vector LIMIT {K}", qvec
            )
            latencies.append((time.perf_counter() - t0) * 1000)
            predicted.append([r["id"] for r in rows])

        recall = _recall(ground_truth, predicted)
        p50, p95, p99, qps = _trimmed_percentiles(latencies, WARMUP_FRACTION)

        return VariantResult(
            variant, description + f" [{index_note}]", available=True,
            recall_at_10=recall, p50_ms=p50, p95_ms=p95, p99_ms=p99, qps=qps,
        )
    finally:
        await conn.execute(f"DROP TABLE IF EXISTS {table}")
        await conn.close()


# ----------------------------------------------------------------------------
# Варианты 2-4: Qdrant (plain / quantized / "как в production" с кэшем)
# ----------------------------------------------------------------------------

async def _qdrant_variant(
    host: str,
    grpc_port: int,
    http_port: int,
    vectors: np.ndarray,
    queries: np.ndarray,
    ground_truth: list[list[int]],
    *,
    variant: str,
    description: str,
    enable_quantization: bool,
    simulate_cache: bool,
) -> VariantResult:
    try:
        from qdrant_client import AsyncQdrantClient
        from qdrant_client.http import models
    except ImportError:
        return VariantResult(variant, description, available=False, notes="qdrant-client не установлен")

    try:
        client = AsyncQdrantClient(host=host, port=http_port, grpc_port=grpc_port, prefer_grpc=True, timeout=10)
        await client.get_collections()
    except Exception as exc:
        return VariantResult(variant, description, available=False, notes=f"Qdrant недоступен: {exc}")

    name = f"bench_{uuid.uuid4().hex[:10]}"
    try:
        quant_cfg = None
        if enable_quantization:
            quant_cfg = models.ScalarQuantization(
                scalar=models.ScalarQuantizationConfig(
                    type=models.ScalarType.INT8, quantile=0.99, always_ram=True,
                )
            )

        await client.create_collection(
            collection_name=name,
            vectors_config=models.VectorParams(size=vectors.shape[1], distance=models.Distance.COSINE),
            hnsw_config=models.HnswConfigDiff(m=HNSW_M, ef_construct=HNSW_EF_CONSTRUCT),
            quantization_config=quant_cfg,
            optimizers_config=models.OptimizersConfigDiff(indexing_threshold=0),
        )

        logger.info("[%s] Inserting %d vectors", variant, len(vectors))
        batch = 500
        for i in range(0, len(vectors), batch):
            chunk = vectors[i : i + batch]
            await client.upsert(
                collection_name=name,
                points=[
                    models.PointStruct(id=i + j, vector=chunk[j].tolist())
                    for j in range(len(chunk))
                ],
                wait=False,
            )
        await client.update_collection(
            collection_name=name,
            optimizers_config=models.OptimizersConfigDiff(indexing_threshold=1),
        )
        while True:
            info = await client.get_collection(name)
            if info.status == models.CollectionStatus.GREEN:
                break
            await asyncio.sleep(0.5)

        search_params = models.SearchParams(
            hnsw_ef=HNSW_EF_SEARCH,
            quantization=models.QuantizationSearchParams(rescore=True) if enable_quantization else None,
        )

        # Простейшая in-memory кэш-эмуляция того же вида, что и SearchCache в проде:
        # ключ — эмбеддинг запроса (у нас запросы уникальны в рамках прогона,
        # так что "тёплого" повторного вызова не будет естественным путём —
        # эмулируем его явно: часть запросов дублируем, чтобы измерить cache-hit latency).
        cold_latencies: list[float] = []
        warm_latencies: list[float] = []
        predicted: list[list[int]] = []
        local_cache: dict[int, list[int]] = {}

        for idx, q in enumerate(queries):
            simulate_repeat = simulate_cache and idx % 3 == 0  # каждый 3-й запрос — "повтор"
            cache_key = idx if not simulate_repeat else idx - (idx % 3)

            if simulate_cache and cache_key in local_cache:
                t0 = time.perf_counter()
                result_ids = local_cache[cache_key]  # dict lookup — то же самое, что делает SearchCache
                warm_latencies.append((time.perf_counter() - t0) * 1000)
                predicted.append(result_ids)
                continue

            t0 = time.perf_counter()
            res = await client.query_points(
                collection_name=name,
                query=q.tolist(),
                limit=K,
                search_params=search_params,
            )
            cold_latencies.append((time.perf_counter() - t0) * 1000)
            result_ids = [p.id for p in res.points]
            predicted.append(result_ids)
            if simulate_cache:
                local_cache[idx] = result_ids

        recall = _recall(ground_truth, predicted)

        if simulate_cache and warm_latencies:
            # Блендим cold/warm по ASSUMED_CACHE_HIT_RATE, а не по фактической доле
            # в этом прогоне (она зависит от искусственного шаблона повторов выше) —
            # честно показываем это как предположение о реальном трафике.
            all_latencies = (
                cold_latencies * int(round((1 - ASSUMED_CACHE_HIT_RATE) * 100))
                + warm_latencies * int(round(ASSUMED_CACHE_HIT_RATE * 100))
            )
            p50, p95, p99, qps = _trimmed_percentiles(sorted(all_latencies), 0.0)
            notes = (
                f"Блендинг cold/warm при предполагаемом cache_hit_rate={ASSUMED_CACHE_HIT_RATE:.0%} "
                f"(assumption, замерить реальный на проде через /health). "
                f"Cold p95={_trimmed_percentiles(cold_latencies, WARMUP_FRACTION)[1]:.2f}ms, "
                f"warm p95={_trimmed_percentiles(warm_latencies, 0.0)[1] if len(warm_latencies) > 1 else warm_latencies[0]:.3f}ms"
            )
        else:
            p50, p95, p99, qps = _trimmed_percentiles(cold_latencies, WARMUP_FRACTION)
            notes = ""

        return VariantResult(
            variant, description, available=True,
            recall_at_10=recall, p50_ms=p50, p95_ms=p95, p99_ms=p99, qps=qps, notes=notes,
        )
    finally:
        try:
            await client.delete_collection(name)
        except Exception:
            pass
        await client.close()


async def run_all(
    database_url: str,
    qdrant_host: str,
    qdrant_grpc_port: int,
    qdrant_http_port: int,
    n_vectors: int,
    n_queries: int,
) -> list[VariantResult]:
    vectors = _gen_vectors(n_vectors, DIM, seed=42)
    queries = _gen_vectors(n_queries, DIM, seed=123)
    ground_truth = _brute_force_ground_truth(vectors, queries, K)

    results: list[VariantResult] = []

    results.append(await bench_postgres_pgvector(database_url, vectors, queries, ground_truth))

    results.append(await _qdrant_variant(
        qdrant_host, qdrant_grpc_port, qdrant_http_port, vectors, queries, ground_truth,
        variant="2. Qdrant (без квантования)",
        description=f"Qdrant HNSW (m={HNSW_M}, ef_construct={HNSW_EF_CONSTRUCT}), ef_search={HNSW_EF_SEARCH}, без квантования, без кэша.",
        enable_quantization=False, simulate_cache=False,
    ))

    results.append(await _qdrant_variant(
        qdrant_host, qdrant_grpc_port, qdrant_http_port, vectors, queries, ground_truth,
        variant="3. Qdrant + INT8 квантование",
        description=f"Qdrant HNSW (m={HNSW_M}, ef_construct={HNSW_EF_CONSTRUCT}) + scalar INT8 quantization (quantile=0.99, rescore=True), без кэша.",
        enable_quantization=True, simulate_cache=False,
    ))

    results.append(await _qdrant_variant(
        qdrant_host, qdrant_grpc_port, qdrant_http_port, vectors, queries, ground_truth,
        variant="4. Полная инфраструктура проекта",
        description=(
            f"Qdrant HNSW (m={HNSW_M}, ef_construct={HNSW_EF_CONSTRUCT}) + INT8 квантование "
            f"+ слой кэширования результатов (как в SearchCache проекта). Именно эта конфигурация "
            f"развёрнута в продакшене (см. collection_service.py, config.py)."
        ),
        enable_quantization=True, simulate_cache=True,
    ))

    return results


# ----------------------------------------------------------------------------
# Отчёты: markdown (быстрый просмотр) + docx (для демонстрации заказчику)
# ----------------------------------------------------------------------------

def _pick_recommendation(results: list[VariantResult], recall_threshold: float = 0.95) -> str:
    candidates = [r for r in results if r.available and r.recall_at_10 is not None and r.recall_at_10 >= recall_threshold]
    if not candidates:
        return (
            "Ни один из доступных вариантов не достиг recall@10 >= "
            f"{recall_threshold:.0%} в этом прогоне — увеличьте ef_search или проверьте настройки индекса "
            "перед принятием решения."
        )
    best = min(candidates, key=lambda r: r.p95_ms)
    return (
        f"Наиболее эффективный вариант при recall@10 >= {recall_threshold:.0%}: **{best.variant}** "
        f"— p95 latency {best.p95_ms:.2f} ms, throughput {best.qps:.1f} QPS, recall@10 {best.recall_at_10:.2%}."
    )


def write_markdown_report(results: list[VariantResult], path: Path, n_vectors: int, n_queries: int) -> None:
    lines = [
        "# Сравнение методов векторного поиска",
        "",
        f"Датасет: {n_vectors:,} векторов, размерность {DIM}, {n_queries} тестовых запросов, метрика — cosine.".replace(",", " "),
        f"K=10, ground truth — brute-force (numpy), одинаковый для всех вариантов.",
        "",
        "| Вариант | Доступен | Recall@10 | p50 (ms) | p95 (ms) | p99 (ms) | QPS |",
        "|---|---|---|---|---|---|---|",
    ]
    for r in results:
        if not r.available:
            lines.append(f"| {r.variant} | ❌ ({r.notes}) | — | — | — | — | — |")
            continue
        lines.append(
            f"| {r.variant} | ✅ | {r.recall_at_10:.2%} | {r.p50_ms:.2f} | {r.p95_ms:.2f} | {r.p99_ms:.2f} | {r.qps:.1f} |"
        )
    lines += ["", "## Описание вариантов", ""]
    for r in results:
        lines.append(f"**{r.variant}** — {r.description}")
        if r.notes:
            lines.append(f"  _Заметка: {r.notes}_")
        lines.append("")
    lines += ["## Рекомендация", "", _pick_recommendation(results)]
    path.write_text("\n".join(lines), encoding="utf-8")


def write_docx_report(results: list[VariantResult], path: Path, n_vectors: int, n_queries: int) -> None:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    doc = Document()

    title = doc.add_heading("Исследование: сравнение методов векторного поиска", level=0)

    doc.add_paragraph(
        f"Датасет: {n_vectors:,} векторов, размерность {DIM}, {n_queries} тестовых запросов. "
        f"Метрика — косинусное сходство. K=10. Ground truth (для расчёта recall) — "
        f"brute-force вычисление в numpy, единое для всех вариантов, независимое от какой-либо БД."
        .replace(",", " ")
    )

    doc.add_heading("Методология", level=1)
    doc.add_paragraph(
        "Один и тот же набор векторов и запросов прогонялся через 4 конфигурации поисковой "
        "инфраструктуры. Каждый вариант оценивался по трём осям: точность (recall@10 относительно "
        "точного brute-force поиска), скорость (латентность p50/p95/p99) и throughput (QPS)."
    )
    for r in results:
        p = doc.add_paragraph()
        p.add_run(f"{r.variant}: ").bold = True
        p.add_run(r.description)

    doc.add_heading("Результаты", level=1)

    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, text in enumerate(["Вариант", "Recall@10", "p50 (ms)", "p95 (ms)", "p99 (ms)", "QPS"]):
        hdr[i].text = text
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True

    for r in results:
        row = table.add_row().cells
        row[0].text = r.variant
        if r.available:
            row[1].text = f"{r.recall_at_10:.2%}"
            row[2].text = f"{r.p50_ms:.2f}"
            row[3].text = f"{r.p95_ms:.2f}"
            row[4].text = f"{r.p99_ms:.2f}"
            row[5].text = f"{r.qps:.1f}"
        else:
            for i in range(1, 6):
                row[i].text = "—"
            row[1].text = f"недоступен ({r.notes})"

    available = [r for r in results if r.available]
    if available:
        doc.add_heading("Латентность (p95), мс — ниже лучше", level=1)
        fig, ax = plt.subplots(figsize=(6.5, 3.2))
        names = [r.variant.split(". ", 1)[-1] for r in available]
        ax.bar(names, [r.p95_ms for r in available], color="#2b6cb0")
        ax.set_ylabel("p95, ms")
        plt.xticks(rotation=20, ha="right", fontsize=8)
        plt.tight_layout()
        chart_path = path.with_suffix(".latency.png")
        fig.savefig(chart_path, dpi=150)
        plt.close(fig)
        doc.add_picture(str(chart_path), width=Inches(6))

        doc.add_heading("Recall@10 — выше лучше", level=1)
        fig, ax = plt.subplots(figsize=(6.5, 3.2))
        ax.bar(names, [r.recall_at_10 for r in available], color="#2f855a")
        ax.set_ylabel("recall@10")
        ax.set_ylim(0, 1.05)
        plt.xticks(rotation=20, ha="right", fontsize=8)
        plt.tight_layout()
        chart_path2 = path.with_suffix(".recall.png")
        fig.savefig(chart_path2, dpi=150)
        plt.close(fig)
        doc.add_picture(str(chart_path2), width=Inches(6))

    doc.add_heading("Примечания к варианту 4 (полная инфраструктура)", level=1)
    doc.add_paragraph(
        f"Латентность варианта 4 — это смесь \"холодных\" запросов (реальный поход в Qdrant) и "
        f"\"тёплых\" (попадание в кэш результатов) при предполагаемой доле повторных запросов "
        f"{ASSUMED_CACHE_HIT_RATE:.0%}. Это предположение о характере трафика, а не измеренная величина — "
        f"после недели реальной эксплуатации нужно взять фактический cache_hit_rate из /health "
        f"и пересчитать эту цифру."
    )

    doc.add_heading("Рекомендация", level=1)
    doc.add_paragraph(_pick_recommendation(results))

    doc.save(str(path))


async def main() -> None:
    parser = argparse.ArgumentParser(description="Сравнение методов векторного поиска: Postgres/pgvector vs Qdrant")
    parser.add_argument("--database-url", default="postgresql://postgres:postgres@localhost:5432/products")
    parser.add_argument("--qdrant-host", default="localhost")
    parser.add_argument("--qdrant-grpc-port", type=int, default=6334)
    parser.add_argument("--qdrant-http-port", type=int, default=6333)
    parser.add_argument("--vectors", type=int, default=20_000)
    parser.add_argument("--queries", type=int, default=300)
    parser.add_argument("--output-prefix", default="search_comparison_report")
    args = parser.parse_args()

    results = await run_all(
        args.database_url, args.qdrant_host, args.qdrant_grpc_port, args.qdrant_http_port,
        args.vectors, args.queries,
    )

    json_path = Path(f"{args.output_prefix}.json")
    md_path = Path(f"{args.output_prefix}.md")
    docx_path = Path(f"{args.output_prefix}.docx")

    json_path.write_text(json.dumps([asdict(r) for r in results], ensure_ascii=False, indent=2), encoding="utf-8")
    write_markdown_report(results, md_path, args.vectors, args.queries)
    write_docx_report(results, docx_path, args.vectors, args.queries)

    for r in results:
        status = "OK" if r.available else f"UNAVAILABLE ({r.notes})"
        logger.info("%s: %s", r.variant, status)

    logger.info("Отчёты сохранены: %s, %s, %s", json_path, md_path, docx_path)


if __name__ == "__main__":
    asyncio.run(main())
